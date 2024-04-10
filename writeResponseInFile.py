from docx import Document
import streamlit as st
import os


init = st.session_state



# Adds a row with user input and model response to an existing table in a 
# Word document to document dialogs or interactions with the language model.
def add_row_to_table(doc, user_prompt, fullResponse):
    if not doc.tables:
        
        doc.add_table(rows=1, cols=2)
    table = doc.tables[0]

    heading_cells = table.rows[0].cells
    heading_cells[0].text = 'Input'
    heading_cells[1].text = 'Response'

    new_row = table.add_row()
    new_row.cells[0].text = user_prompt
    new_row.cells[1].text = fullResponse




# Writes the user input and the response of the language model to a Word file and saves it. 
def writeLLMAnswerToFileIfEnabled(file_path, user_prompt, fullResponse, enabled=False):

    try:
        if user_prompt and enabled:

            if os.path.exists(file_path):
                doc = Document(file_path)
            else:
                doc = Document()

            add_row_to_table(doc, init.user_prompt, fullResponse)

            doc.save(file_path)

    except PermissionError:
        st.error("Close .docx file!")